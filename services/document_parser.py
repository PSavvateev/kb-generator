"""
Document Parser - Unified Module

Extracts text, tables, and metadata from PDF, DOCX, and TXT files

Features:
- PDF, DOCX, TXT support
- Table extraction with markdown formatting
- Configurable table handling
- Type-safe with dataclasses
- Comprehensive error handling

Dependencies:
    PyPDF2, pdfplumber, python-docx, chardet
"""

import logging
from pathlib import Path
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field, asdict

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Check dependencies
try:
    import PyPDF2
    import pdfplumber
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

try:
    import chardet
    HAS_CHARDET = True
except ImportError:
    HAS_CHARDET = False


@dataclass
class ExtractedTable:
    """Represents an extracted table"""
    data: List[List[str]]
    markdown: str
    page: Optional[int] = None
    location: Optional[str] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            'data': self.data,
            'markdown': self.markdown,
            'page': self.page,
            'location': self.location
        }


@dataclass
class DocumentMetadata:
    """Document metadata"""
    file_type: str
    file_size_bytes: int
    word_count: int
    char_count: int
    
    title: Optional[str] = None
    author: Optional[str] = None
    subject: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    
    page_count: Optional[int] = None
    has_images: Optional[bool] = None
    paragraph_count: Optional[int] = None
    table_count: Optional[int] = None
    encoding: Optional[str] = None
    line_count: Optional[int] = None
    
    tables: List[Dict[str, Any]] = field(default_factory=list)
    
    def to_dict(self) -> Dict[str, Any]:
        return {k: v for k, v in asdict(self).items() if v is not None}


class TableExtractor:
    """Table extraction utilities"""
    
    @staticmethod
    def table_to_markdown(table_data: List[List[str]], has_header: bool = True) -> str:
        if not table_data or not table_data[0]:
            return ""
        
        def clean_cell(cell):
            if cell is None:
                return ""
            return str(cell).strip().replace('\n', ' ')
        
        cleaned = [[clean_cell(cell) for cell in row] for row in table_data]
        num_cols = len(cleaned[0])
        col_widths = [max(len(row[i]) if i < len(row) else 0 for row in cleaned) for i in range(num_cols)]
        col_widths = [max(w, 3) for w in col_widths]
        
        lines = []
        if has_header and cleaned:
            header = cleaned[0]
            lines.append("| " + " | ".join(header[i].ljust(col_widths[i]) if i < len(header) else " " * col_widths[i] for i in range(num_cols)) + " |")
            lines.append("| " + " | ".join("-" * w for w in col_widths) + " |")
            for row in cleaned[1:]:
                lines.append("| " + " | ".join(row[i].ljust(col_widths[i]) if i < len(row) else " " * col_widths[i] for i in range(num_cols)) + " |")
        else:
            for row in cleaned:
                lines.append("| " + " | ".join(row[i].ljust(col_widths[i]) if i < len(row) else " " * col_widths[i] for i in range(num_cols)) + " |")
        
        return "\n".join(lines)
    
    @staticmethod
    def is_valid_table(table_data: List[List[str]], 
                    min_rows: int = 2, 
                    min_cols: int = 2,
                    strict: bool = True) -> bool:
        """
        Validate if extracted data is a proper table
        
        Args:
            table_data: Raw table data
            min_rows: Minimum number of rows
            min_cols: Minimum number of columns
            strict: Enable strict validation (filters malformed tables)
            
        Returns:
            True if valid table, False otherwise
        """
        # Basic checks
        if not table_data or len(table_data) < min_rows:
            logger.debug(f"Table rejected: insufficient rows ({len(table_data) if table_data else 0} < {min_rows})")
            return False
        
        # Check if has minimum columns
        has_cols = any(len(row) >= min_cols for row in table_data)
        if not has_cols:
            logger.debug("Table rejected: insufficient columns")
            return False
        
        # Check if has actual content
        has_content = any(any(cell and str(cell).strip() for cell in row) for row in table_data)
        if not has_content:
            logger.debug("Table rejected: no content")
            return False
        
        if not strict:
            return True
        
        # STRICT VALIDATION - Filter malformed tables
        
        # 1. Check for too many empty/None cells (>60% empty = not a real table)
        total_cells = sum(len(row) for row in table_data)
        empty_cells = sum(
            1 for row in table_data 
            for cell in row 
            if not cell or (isinstance(cell, str) and not cell.strip())
        )
        
        if total_cells > 0:
            empty_ratio = empty_cells / total_cells
            if empty_ratio > 0.6:
                logger.debug(f"Table rejected: too many empty cells ({empty_ratio:.1%})")
                return False
        
        # 2. Check if first column is completely empty (common in PDF extraction errors)
        if len(table_data[0]) > 1:  # Only check if table has multiple columns
            first_col_values = [row[0] if row else None for row in table_data]
            non_empty_first_col = [
                v for v in first_col_values 
                if v and (not isinstance(v, str) or v.strip())
            ]
            
            if len(non_empty_first_col) == 0:
                logger.debug("Table rejected: first column completely empty")
                return False
        
        # 3. Check if it's a single-cell table (not a real table)
        if len(table_data) == 1 and len(table_data[0]) <= 1:
            logger.debug("Table rejected: single cell")
            return False
        
        # 4. Check for "fake tables" - single row/column with huge text blob
        # These are usually visual boxes misidentified as tables
        if len(table_data) <= 2 and len(table_data[0]) <= 2:
            # Check if any cell has >500 characters (likely a text block, not tabular data)
            for row in table_data:
                for cell in row:
                    if cell and len(str(cell)) > 500:
                        logger.debug("Table rejected: contains large text blob (likely not a table)")
                        return False
        
        # 5. Check column consistency - real tables have consistent column counts
        col_counts = [len(row) for row in table_data]
        if len(set(col_counts)) > len(table_data) * 0.3:  # >30% of rows have different column counts
            logger.debug("Table rejected: inconsistent column structure")
            return False
        
        logger.debug(f"Table validated: {len(table_data)} rows, ~{len(table_data[0])} cols")
        return True


class DocumentParser:
    """Parse documents and extract text, tables, and metadata"""
    
    HAS_PDF_SUPPORT = HAS_PDF_SUPPORT
    HAS_DOCX_SUPPORT = HAS_DOCX_SUPPORT
    HAS_CHARDET = HAS_CHARDET
    DEFAULT_MAX_FILE_SIZE = 100 * 1024 * 1024
    
    def __init__(self, 
                 max_file_size: int = DEFAULT_MAX_FILE_SIZE,
                 extract_tables: bool = True,
                 tables_as_markdown: bool = True,
                 strict_table_validation: bool = True):
        """
        Args:
            max_file_size: Maximum file size in bytes
            extract_tables: Extract tables from PDFs/DOCX
            tables_as_markdown: Include tables in text as markdown
        """
        self.supported_formats = {'.pdf', '.docx', '.txt'}
        self.max_file_size = max_file_size
        self.extract_tables = extract_tables
        self.tables_as_markdown = tables_as_markdown
        self.strict_table_validation = strict_table_validation
        self.table_extractor = TableExtractor()
    
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse document
        
        Returns:
            {
                'text': str,
                'metadata': dict,
                'tables': list[dict]
            }
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {file_ext}")
        
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size:
            raise ValueError(f"File too large: {file_size / (1024*1024):.1f}MB")
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path, file_size)
        elif file_ext == '.docx':
            return self._parse_docx(file_path, file_size)
        elif file_ext == '.txt':
            return self._parse_txt(file_path, file_size)
    
    def _parse_pdf(self, file_path: Path, file_size: int) -> Dict[str, Any]:
        if not self.HAS_PDF_SUPPORT:
            raise ImportError("Install: pip install PyPDF2 pdfplumber")
        
        import pdfplumber
        
        text_parts = []
        tables = []
        page_count = 0
        has_images = False
        
        try:
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    
                    if self.extract_tables:
                        page_tables = page.extract_tables()
                        if page_tables:
                            for table_idx, table_data in enumerate(page_tables):
                                if self.table_extractor.is_valid_table(table_data, strict=self.strict_table_validation):
                                    markdown = self.table_extractor.table_to_markdown(table_data)
                                    
                                    extracted_table = ExtractedTable(
                                        data=table_data,
                                        markdown=markdown,
                                        page=page_num,
                                        location=f"page {page_num}, table {table_idx + 1}"
                                    )
                                    tables.append(extracted_table)
                                    
                                    if self.tables_as_markdown:
                                        text_parts.append(f"\n\n{markdown}\n")
                    
                    if page_text:
                        text_parts.append(page_text)
                    
                    if page.images:
                        has_images = True
        
        except Exception as e:
            logger.warning(f"pdfplumber failed, trying PyPDF2: {e}")
            return self._parse_pdf_pypdf2(file_path, file_size)
        
        text = "\n\n".join(text_parts)
        if not text.strip():
            raise ValueError("No text extracted")
        
        metadata = DocumentMetadata(
            file_type="pdf",
            file_size_bytes=file_size,
            word_count=len(text.split()),
            char_count=len(text),
            page_count=page_count,
            has_images=has_images,
            table_count=len(tables),
            tables=[t.to_dict() for t in tables]
        )
        
        return {
            "text": text,
            "metadata": metadata.to_dict(),
            "tables": [t.to_dict() for t in tables]
        }
    
    def _parse_pdf_pypdf2(self, file_path: Path, file_size: int) -> Dict[str, Any]:
        import PyPDF2
        
        text_parts = []
        with open(file_path, 'rb') as file:
            pdf = PyPDF2.PdfReader(file)
            page_count = len(pdf.pages)
            for page in pdf.pages:
                if page_text := page.extract_text():
                    text_parts.append(page_text)
        
        text = "\n\n".join(text_parts)
        if not text.strip():
            raise ValueError("No text extracted")
        
        metadata = DocumentMetadata(
            file_type="pdf",
            file_size_bytes=file_size,
            word_count=len(text.split()),
            char_count=len(text),
            page_count=page_count,
            table_count=0
        )
        
        return {"text": text, "metadata": metadata.to_dict(), "tables": []}
    
    def _parse_docx(self, file_path: Path, file_size: int) -> Dict[str, Any]:
        if not self.HAS_DOCX_SUPPORT:
            raise ImportError("Install: pip install python-docx")
        
        from docx import Document as DocxDocument
        
        doc = DocxDocument(file_path)
        text_parts = []
        tables = []
        
        for element in doc.element.body:
            if element.tag.endswith('p'):
                for para in doc.paragraphs:
                    if para._element == element:
                        if para.text.strip():
                            text_parts.append(para.text)
                        break
            
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                        
                        if self.extract_tables and self.table_extractor.is_valid_table(table_data, strict=self.strict_table_validation):
                            markdown = self.table_extractor.table_to_markdown(table_data)
                            
                            extracted_table = ExtractedTable(
                                data=table_data,
                                markdown=markdown,
                                location=f"table {len(tables) + 1}"
                            )
                            tables.append(extracted_table)
                            
                            if self.tables_as_markdown:
                                text_parts.append(f"\n\n{markdown}\n")
                        break
        
        text = "\n\n".join(text_parts)
        
        props = doc.core_properties if hasattr(doc, 'core_properties') else None
        metadata = DocumentMetadata(
            file_type="docx",
            file_size_bytes=file_size,
            word_count=len(text.split()) if text else 0,
            char_count=len(text),
            paragraph_count=len(doc.paragraphs),
            table_count=len(tables),
            tables=[t.to_dict() for t in tables],
            title=props.title if props and props.title else None,
            author=props.author if props and props.author else None,
            subject=props.subject if props and props.subject else None,
            created_date=props.created.isoformat() if props and props.created else None,
            modified_date=props.modified.isoformat() if props and props.modified else None
        )
        
        return {"text": text, "metadata": metadata.to_dict(), "tables": [t.to_dict() for t in tables]}
    
    def _parse_txt(self, file_path: Path, file_size: int) -> Dict[str, Any]:
        encoding = 'utf-8'
        if self.HAS_CHARDET:
            with open(file_path, 'rb') as f:
                result = chardet.detect(f.read())
                if result['encoding']:
                    encoding = result['encoding']
        
        with open(file_path, 'r', encoding=encoding, errors='replace') as f:
            text = f.read()
        
        metadata = DocumentMetadata(
            file_type="txt",
            file_size_bytes=file_size,
            word_count=len(text.split()) if text else 0,
            char_count=len(text),
            line_count=len(text.split('\n')),
            encoding=encoding,
            table_count=0
        )
        
        return {"text": text, "metadata": metadata.to_dict(), "tables": []}


def parse_document(file_path: str, 
                  max_file_size: int = 100 * 1024 * 1024,
                  extract_tables: bool = True,
                  tables_as_markdown: bool = True,
                  strict_table_validation: bool = True) -> Dict[str, Any]:
    """Convenience function to parse a document"""
    parser = DocumentParser(
        max_file_size=max_file_size,
        extract_tables=extract_tables,
        tables_as_markdown=tables_as_markdown,
        strict_table_validation=strict_table_validation
    )
    return parser.parse(file_path)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) > 1:
        result = parse_document(sys.argv[1])
        print(f"Type: {result['metadata']['file_type']}")
        print(f"Words: {result['metadata']['word_count']}")
        print(f"Tables: {result['metadata']['table_count']}")
        if result['tables']:
            for i, table in enumerate(result['tables'], 1):
                print(f"\nTable {i}: {table.get('location')}")
    else:
        print("Usage: python document_parser.py <file>")